"""Konverterar nedladdade Drive-filer (PDF/DOCX/ODT/text) till klartext.

Använd `extract_text(path, mime_type)`. Höjer `UnsupportedFormatError`
för format som inte stöds — uppströms ska det rapporteras som "skipped"
med en tydlig anledning.

Importerna är lazy så att verktyget kan startas även om någon enskild
konverter saknas; bara när den filtypen faktiskt dyker upp visas felet.
"""

from __future__ import annotations

from pathlib import Path


class UnsupportedFormatError(RuntimeError):
    pass


# MIME-typer Drive använder för uppladdade filer
PDF_MIMES = {"application/pdf"}
DOCX_MIMES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
ODT_MIMES = {"application/vnd.oasis.opendocument.text"}
TEXT_MIMES = {"text/plain", "text/markdown", "application/x-markdown"}
DOC_LEGACY_MIMES = {"application/msword"}  # .doc — saknar bra Python-stöd
RTF_MIMES = {"application/rtf", "text/rtf"}


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise UnsupportedFormatError(
            "pypdf saknas — kör `pip install pypdf` i venv"
        ) from e

    reader = PdfReader(str(path))
    chunks: list[str] = []
    for page in reader.pages:
        try:
            chunks.append(page.extract_text() or "")
        except Exception as e:
            chunks.append(f"\n[fel vid extrahering av sida: {e}]\n")
    return "\n\n".join(c.strip() for c in chunks if c.strip())


def _read_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as e:
        raise UnsupportedFormatError(
            "python-docx saknas — kör `pip install python-docx` i venv"
        ) from e

    doc = Document(str(path))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # Tabeller kan innehålla viktig text (t.ex. arbetsblad-strukturer).
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return "\n".join(parts)


def _read_odt(path: Path) -> str:
    try:
        from odf import text as odf_text, teletype
        from odf.opendocument import load
    except ImportError as e:
        raise UnsupportedFormatError(
            "odfpy saknas — kör `pip install odfpy` i venv"
        ) from e

    doc = load(str(path))
    parts: list[str] = []
    for paragraph in doc.getElementsByType(odf_text.P):
        s = teletype.extractText(paragraph)
        if s.strip():
            parts.append(s)
    for heading in doc.getElementsByType(odf_text.H):
        s = teletype.extractText(heading)
        if s.strip():
            parts.append(s)
    return "\n".join(parts)


def _read_plaintext(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def extract_text(path: Path, mime_type: str | None) -> str:
    """Returnerar klartext för en fil. Höjer UnsupportedFormatError vid
    okänt format eller saknad konverter."""
    mt = (mime_type or "").lower()

    if mt in PDF_MIMES:
        return _read_pdf(path)
    if mt in DOCX_MIMES:
        return _read_docx(path)
    if mt in ODT_MIMES:
        return _read_odt(path)
    if mt in TEXT_MIMES:
        return _read_plaintext(path)
    if mt in DOC_LEGACY_MIMES:
        raise UnsupportedFormatError(
            "Legacy .doc stöds inte — be eleven spara om som .docx eller PDF"
        )
    if mt in RTF_MIMES:
        raise UnsupportedFormatError(
            "RTF stöds inte — be eleven spara om som .docx, PDF eller Google Doc"
        )

    # Fall back på filändelse om mime saknas/är generisk
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix == ".docx":
        return _read_docx(path)
    if suffix == ".odt":
        return _read_odt(path)
    if suffix in {".txt", ".md"}:
        return _read_plaintext(path)

    raise UnsupportedFormatError(
        f"Okänt filformat: mime={mime_type!r} suffix={suffix!r}"
    )
